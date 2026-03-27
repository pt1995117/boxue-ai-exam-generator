from docx import Document

from generate_knowledge_slices import (
    group_and_slice,
    merge_short_slices_at_parent,
    process_document,
    repair_flattened_paths,
)


def test_repair_flattened_paths_inserts_missing_l4_parent():
    slices = [
        {"完整路径": "第一篇 > 第一章 > 第一节 > 一、总述"},
        {"完整路径": "第一篇 > 第一章 > 第一节 > 三、房地产市场周期波动"},
        {"完整路径": "第一篇 > 第一章 > 第一节 > （一）周期影响因素"},
        {"完整路径": "第一篇 > 第一章 > 第一节 > （二）市场分析指标"},
    ]

    fixed = repair_flattened_paths(slices)

    assert fixed[2]["完整路径"] == "第一篇 > 第一章 > 第一节 > 三、房地产市场周期波动 > （一）周期影响因素"
    assert fixed[3]["完整路径"] == "第一篇 > 第一章 > 第一节 > 三、房地产市场周期波动 > （二）市场分析指标"


def test_repair_flattened_paths_keeps_already_full_paths():
    slices = [
        {"完整路径": "第一篇 > 第一章 > 第一节 > 三、房地产市场周期波动 > （一）周期影响因素"},
    ]

    fixed = repair_flattened_paths(slices)

    assert fixed[0]["完整路径"] == "第一篇 > 第一章 > 第一节 > 三、房地产市场周期波动 > （一）周期影响因素"


def test_group_and_slice_does_not_leak_internal_state_field():
    elements = [
        {
            "type": "heading",
            "level": 4,
            "text": "一、总述",
            "path": "第一篇 > 第一章 > 第一节 > 一、总述",
            "mastery": "了解",
        },
        {
            "type": "paragraph",
            "text": "这是正文",
            "path": "第一篇 > 第一章 > 第一节 > 一、总述",
            "mastery": "了解",
            "images": [],
        },
    ]

    out = group_and_slice(elements, api_key="")

    assert len(out) == 1
    assert "_in_example" not in out[0]


def test_merge_short_slices_keeps_route_hint_in_merged_content():
    slices = [
        {
            "完整路径": "第一篇 > 第一章",
            "掌握程度": "了解",
            "结构化内容": {
                "key_params": [],
                "rules": [],
                "context_before": "",
                "tables": [],
                "context_after": "",
                "images": [],
                "image_anchors": [],
                "formulas": [],
                "examples": [],
            },
            "metadata": {"类型": "自动组装"},
        },
        {
            "完整路径": "第一篇 > 第一章 > 第一节",
            "掌握程度": "了解",
            "结构化内容": {
                "key_params": [],
                "rules": [],
                "context_before": "正文A",
                "tables": [],
                "context_after": "",
                "images": [],
                "image_anchors": [],
                "formulas": [],
                "examples": [],
            },
            "metadata": {"类型": "自动组装"},
        },
    ]

    merged = merge_short_slices_at_parent(slices, min_chars=200)
    merged_text = merged[0]["结构化内容"]["context_before"]
    assert "【路由】第一节" in merged_text
    assert "正文A" in merged_text


def test_process_document_splits_custom_cover_title_by_style_and_toc(tmp_path, monkeypatch):
    monkeypatch.setenv("SLICE_DISABLE_IMAGE_OCR", "1")

    doc_path = tmp_path / "custom_cover_title.docx"
    doc = Document()
    toc_style = doc.styles.add_style("toc 1", 1)
    toc_style.base_style = doc.styles["Normal"]
    heading1_style = doc.styles.add_style("自定义篇", 1)
    heading1_style.base_style = doc.styles["Heading 1"]

    p = doc.add_paragraph("第一篇 总论导读\t1")
    p.style = toc_style

    p = doc.add_paragraph("第四篇  服务保障")
    p.style = heading1_style
    doc.add_paragraph("第五章  个人信息安全合规")
    doc.add_paragraph("四、违规处理个人信息的法律责任（了解）")
    doc.add_paragraph("上一节正文")
    p = doc.add_paragraph("下册  干部手册")
    p.style = heading1_style
    p = doc.add_paragraph("总论导读")
    p.style = heading1_style
    doc.add_paragraph("总论正文第一段")
    p = doc.add_paragraph("第二篇  干部管理篇")
    p.style = heading1_style
    doc.add_paragraph("干部管理正文")
    doc.save(str(doc_path))

    elements = process_document(str(doc_path), api_key="", extract_dir=str(tmp_path / "images"))
    slices = group_and_slice(elements, api_key="")

    paths = [s["完整路径"] for s in slices]
    assert any(p.replace("\t", " ").strip() == "第一篇 总论导读" for p in paths)
    assert "第二篇  干部管理篇" in paths

    service_slice = next(s for s in slices if s["完整路径"].endswith("四、违规处理个人信息的法律责任"))
    assert "下册  干部手册" not in service_slice["结构化内容"]["context_before"]
    assert "总论导读" not in service_slice["结构化内容"]["context_before"]

    strategy_slice = next(
        s for s in slices if s["完整路径"].replace("\t", " ").strip() == "第一篇 总论导读"
    )
    assert "总论正文第一段" in strategy_slice["结构化内容"]["context_before"]
