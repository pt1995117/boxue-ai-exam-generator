from __future__ import annotations

import json
import re
from pathlib import Path
from typing import Any

import pandas as pd

REFERENCE_COLUMNS = [
    "题干",
    "考点",
    "选项1",
    "选项2",
    "选项3",
    "选项4",
    "选项5",
    "选项6",
    "选项7",
    "选项8",
    "正确答案",
    "解析",
    "难度值",
]

_COLUMN_ALIASES = {
    "题干": "题干",
    "题目": "题干",
    "question": "题干",
    "stem": "题干",
    "题型": "题型",
    "考点": "考点",
    "knowledge": "考点",
    "知识点": "考点",
    "答案": "正确答案",
    "正确答案": "正确答案",
    "answer": "正确答案",
    "解析": "解析",
    "分析": "解析",
    "explanation": "解析",
    "analysis": "解析",
    "难度": "难度值",
    "难度值": "难度值",
    "difficulty": "难度值",
}


def empty_reference_df() -> pd.DataFrame:
    return pd.DataFrame(columns=REFERENCE_COLUMNS)


def _clean_text(value: Any) -> str:
    if value is None:
        return ""
    if isinstance(value, float) and pd.isna(value):
        return ""
    text = str(value).strip()
    if text.lower() == "nan":
        return ""
    return text


def _normalize_column_name(name: Any) -> str:
    raw = _clean_text(name)
    if not raw:
        return ""
    lowered = raw.lower()
    if lowered in _COLUMN_ALIASES:
        return _COLUMN_ALIASES[lowered]
    if raw in _COLUMN_ALIASES:
        return _COLUMN_ALIASES[raw]
    m = re.match(r"^(?:选项|option)\s*([1-8a-hA-H])$", raw, re.IGNORECASE)
    if m:
        token = m.group(1).upper()
        if token.isdigit():
            return f"选项{token}"
        return f"选项{ord(token) - ord('A') + 1}"
    return raw


def _normalize_answer(value: Any) -> str:
    text = _clean_text(value).upper()
    if not text:
        return ""
    text = (
        text.replace("，", "")
        .replace(",", "")
        .replace(" ", "")
        .replace("、", "")
        .replace("；", "")
        .replace(";", "")
    )
    if text in {"正确", "对", "√"}:
        return "A"
    if text in {"错误", "错", "×"}:
        return "B"
    letters = "".join(ch for ch in text if ch in "ABCDEFGH")
    return letters or text


def _append_field(target: dict[str, Any], key: str, value: str) -> None:
    if not value:
        return
    if key == "解析" and target.get(key):
        target[key] = f"{target[key]}\n{value}"
        return
    if key == "题干" and target.get(key):
        target[key] = f"{target[key]} {value}".strip()
        return
    target[key] = value


def normalize_reference_df(df: pd.DataFrame | None) -> pd.DataFrame:
    if df is None:
        return empty_reference_df()
    if df.empty:
        out = df.copy()
    else:
        renamed = {_clean_text(col): _normalize_column_name(col) for col in df.columns}
        out = df.rename(columns=renamed).copy()
    for col in REFERENCE_COLUMNS:
        if col not in out.columns:
            out[col] = ""
    for col in REFERENCE_COLUMNS:
        out[col] = out[col].map(_clean_text)
    out["正确答案"] = out["正确答案"].map(_normalize_answer)
    valid_mask = out["题干"].astype(str).str.strip() != ""
    return out.loc[valid_mask, REFERENCE_COLUMNS].reset_index(drop=True)


def _rows_to_df(rows: list[dict[str, Any]]) -> pd.DataFrame:
    if not rows:
        return empty_reference_df()
    return normalize_reference_df(pd.DataFrame(rows))


def _parse_docx_tables(path: Path) -> list[dict[str, Any]]:
    from docx import Document

    rows: list[dict[str, Any]] = []
    doc = Document(str(path))
    for table in doc.tables:
        if not table.rows:
            continue
        headers = [_normalize_column_name(cell.text) for cell in table.rows[0].cells]
        if "题干" not in headers:
            continue
        for row in table.rows[1:]:
            item: dict[str, Any] = {}
            for idx, cell in enumerate(row.cells):
                header = headers[idx] if idx < len(headers) else ""
                if not header:
                    continue
                item[header] = _clean_text(cell.text)
            if _clean_text(item.get("题干")):
                rows.append(item)
    return rows


def _question_start(line: str) -> str | None:
    patterns = [
        r"^\s*第\s*\d+\s*题[\.、:：\s]*(.*)$",
        r"^\s*\d+\s*[\.\、\)\）]\s*(.*)$",
        r"^\s*[一二三四五六七八九十]+\s*[\.\、\)\）]\s*(.*)$",
    ]
    for pattern in patterns:
        m = re.match(pattern, line)
        if m:
            return m.group(1).strip()
    return None


def _parse_question_text(text: str) -> list[dict[str, Any]]:
    rows: list[dict[str, Any]] = []
    current: dict[str, Any] | None = None
    last_option_key = ""
    mode = "stem"

    def flush() -> None:
        nonlocal current, last_option_key, mode
        if current and _clean_text(current.get("题干")):
            rows.append(current)
        current = None
        last_option_key = ""
        mode = "stem"

    for raw_line in text.splitlines():
        line = raw_line.strip()
        if not line:
            continue
        starter = _question_start(line)
        if starter is not None:
            flush()
            current = {"题干": starter}
            continue
        if current is None:
            continue

        option_match = re.match(r"^\s*([A-H])[\.\、\)\）:：\s]+(.*)$", line, re.IGNORECASE)
        if option_match:
            letter = option_match.group(1).upper()
            last_option_key = f"选项{ord(letter) - ord('A') + 1}"
            current[last_option_key] = option_match.group(2).strip()
            mode = "option"
            continue

        answer_match = re.match(r"^\s*(?:答案|正确答案)[:：]\s*(.+)$", line, re.IGNORECASE)
        if answer_match:
            current["正确答案"] = _normalize_answer(answer_match.group(1))
            mode = "answer"
            continue

        analysis_match = re.match(r"^\s*(?:解析|分析|解释)[:：]\s*(.*)$", line, re.IGNORECASE)
        if analysis_match:
            current["解析"] = analysis_match.group(1).strip()
            mode = "解析"
            continue

        point_match = re.match(r"^\s*(?:考点|知识点)[:：]\s*(.*)$", line, re.IGNORECASE)
        if point_match:
            current["考点"] = point_match.group(1).strip()
            continue

        if mode == "option" and last_option_key:
            _append_field(current, last_option_key, line)
        elif mode == "解析":
            _append_field(current, "解析", line)
        else:
            _append_field(current, "题干", line)
            mode = "stem"

    flush()
    return rows


def _parse_docx_questions(path: Path) -> pd.DataFrame:
    from docx import Document

    table_rows = _parse_docx_tables(path)
    if table_rows:
        return _rows_to_df(table_rows)
    doc = Document(str(path))
    text = "\n".join(_clean_text(p.text) for p in doc.paragraphs if _clean_text(p.text))
    return _rows_to_df(_parse_question_text(text))


def _parse_text_questions(path: Path) -> pd.DataFrame:
    text = path.read_text(encoding="utf-8", errors="ignore")
    return _rows_to_df(_parse_question_text(text))


def _parse_json_questions(path: Path) -> pd.DataFrame:
    try:
        payload = json.loads(path.read_text(encoding="utf-8"))
    except Exception:
        return empty_reference_df()
    if isinstance(payload, list):
        return normalize_reference_df(pd.DataFrame(payload))
    if isinstance(payload, dict):
        for key in ("items", "questions", "data"):
            value = payload.get(key)
            if isinstance(value, list):
                return normalize_reference_df(pd.DataFrame(value))
        return normalize_reference_df(pd.DataFrame([payload]))
    return empty_reference_df()


def _parse_jsonl_questions(path: Path) -> pd.DataFrame:
    rows: list[dict[str, Any]] = []
    with path.open("r", encoding="utf-8") as f:
        for line in f:
            raw = line.strip()
            if not raw:
                continue
            try:
                item = json.loads(raw)
            except json.JSONDecodeError:
                continue
            if isinstance(item, dict):
                rows.append(item)
    return normalize_reference_df(pd.DataFrame(rows))


def load_reference_questions(path: str | Path) -> pd.DataFrame:
    ref_path = Path(path)
    if not ref_path.exists() or not ref_path.is_file():
        return empty_reference_df()

    suffix = ref_path.suffix.lower()
    try:
        if suffix in {".xlsx", ".xls"}:
            return normalize_reference_df(pd.read_excel(ref_path))
        if suffix == ".csv":
            return normalize_reference_df(pd.read_csv(ref_path))
        if suffix == ".json":
            return _parse_json_questions(ref_path)
        if suffix == ".jsonl":
            return _parse_jsonl_questions(ref_path)
        if suffix == ".docx":
            return _parse_docx_questions(ref_path)
        if suffix in {".txt", ".md"}:
            return _parse_text_questions(ref_path)
    except Exception:
        return empty_reference_df()
    return empty_reference_df()
