import base64

from docx import Document
import pytest

from generate_knowledge_slices import load_config, process_document


_PNG_1X1_BASE64 = (
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8"
    "/x8AAusB9Y9l9XwAAAAASUVORK5CYII="
)


def _write_test_png(path):
    path.write_bytes(base64.b64decode(_PNG_1X1_BASE64))


def _get_image_api_key():
    cfg = load_config()
    provider = (cfg.get("IMAGE_PROVIDER") or "").lower()
    if provider == "ark":
        return (cfg.get("IMAGE_API_KEY") or cfg.get("ARK_API_KEY") or cfg.get("OPENAI_API_KEY") or "").strip()
    return (
        cfg.get("AIT_API_KEY")
        or cfg.get("IMAGE_API_KEY")
        or cfg.get("OPENAI_API_KEY")
        or cfg.get("CRITIC_API_KEY")
        or ""
    ).strip()


def test_process_document_extracts_paragraph_and_table_images(tmp_path, monkeypatch):
    monkeypatch.setenv("SLICE_DISABLE_IMAGE_OCR", "1")

    img_path = tmp_path / "img.png"
    _write_test_png(img_path)

    doc_path = tmp_path / "with_images.docx"
    doc = Document()
    doc.add_heading("第一篇", level=1)

    p = doc.add_paragraph("段落图片：")
    p.add_run().add_picture(str(img_path))

    table = doc.add_table(rows=1, cols=1)
    cell_para = table.cell(0, 0).paragraphs[0]
    cell_para.add_run("单元格图片：")
    cell_para.add_run().add_picture(str(img_path))

    doc.save(str(doc_path))

    elements = process_document(str(doc_path), api_key="", extract_dir=str(tmp_path / "images"))

    para_with_images = [e for e in elements if e.get("type") == "paragraph" and e.get("images")]
    assert para_with_images, "expected at least one paragraph element with extracted images"
    assert para_with_images[0]["images"][0]["analysis"].startswith("(已跳过图片OCR")

    table_elements = [e for e in elements if e.get("type") == "table"]
    assert table_elements, "expected at least one table element"
    assert table_elements[0].get("images"), "expected table-level extracted images"
    anchors = table_elements[0].get("image_anchors") or []
    assert anchors, "expected table image anchors"
    assert anchors[0].get("anchor_type") == "table_cell"


def test_process_document_renders_merged_table_as_html(tmp_path, monkeypatch):
    monkeypatch.setenv("SLICE_DISABLE_IMAGE_OCR", "1")

    doc_path = tmp_path / "merged_table.docx"
    doc = Document()
    doc.add_heading("第一篇", level=1)

    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "表头A"
    table.cell(0, 1).text = "表头B"
    table.cell(1, 0).text = "左侧"
    table.cell(1, 1).text = "右侧"
    table.cell(0, 0).merge(table.cell(0, 1))

    doc.save(str(doc_path))

    elements = process_document(str(doc_path), api_key="", extract_dir=str(tmp_path / "images"))
    table_elements = [e for e in elements if e.get("type") == "table"]

    assert table_elements, "expected at least one table element"
    table_el = table_elements[0]
    assert table_el.get("table_has_merged") is True
    assert table_el.get("table_format") == "html"
    assert "<table>" in table_el.get("text", "")
    assert "colspan=" in table_el.get("text", "")


@pytest.mark.integration
def test_process_document_real_ocr_image_analysis(tmp_path):
    api_key = _get_image_api_key()
    if not api_key:
        pytest.skip("missing image API key in 填写您的Key.txt")

    img_path = tmp_path / "img.png"
    _write_test_png(img_path)

    doc_path = tmp_path / "with_real_ocr.docx"
    doc = Document()
    doc.add_heading("第一篇", level=1)
    p = doc.add_paragraph("段落图片：")
    p.add_run().add_picture(str(img_path))
    doc.save(str(doc_path))

    elements = process_document(str(doc_path), api_key=api_key, extract_dir=str(tmp_path / "images"))

    para_with_images = [e for e in elements if e.get("type") == "paragraph" and e.get("images")]
    assert para_with_images, "expected at least one paragraph element with extracted images"
    analysis = str(para_with_images[0]["images"][0].get("analysis", ""))
    assert analysis and not analysis.startswith("(已跳过图片OCR"), analysis
