#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Generate Knowledge Slices from Word Document
Strictly following 'Final Complete Slice Style Specifications'
"""
import os
import json
import re
import sys
import argparse
import zipfile
import shutil
import html
from datetime import datetime, timezone
from typing import List, Dict, Optional, Tuple
from docx import Document
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.text.paragraph import Paragraph
from tenants_config import tenant_slices_dir
from runtime_paths import load_primary_key_config

# Optional BGE embedding for sub-slicing under level-5 routes
try:
    import numpy as np
except Exception:
    np = None

try:
    from sentence_transformers import SentenceTransformer
except Exception:
    SentenceTransformer = None

# Import helper for image analysis if available
try:
    from process_textbook_images import (
        analyze_image_with_qwen_vl,
        extract_table_from_content,
        normalize_image_analysis_content,
    )
except ImportError:
    analyze_image_with_qwen_vl = None
    extract_table_from_content = None
    normalize_image_analysis_content = None

# Import exam_graph for splitting
try:
    from exam_graph import generate_content
except ImportError:
    sys.path.append(os.path.dirname(os.path.abspath(__file__)))
    try:
        from exam_graph import generate_content
    except ImportError:
        generate_content = None

def load_config():
    return load_primary_key_config()

# --- Image Handling ---

def extract_images_from_docx(docx_path: str, output_dir: str) -> Dict[str, str]:
    """
    Extract all images from docx to output_dir.
    Returns a map: { 'image_part_name': 'local_file_path' }
    """
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)
    
    image_map = {}
    with zipfile.ZipFile(docx_path, 'r') as z:
        for name in z.namelist():
            if name.startswith('word/media/') and not name.endswith('/'):
                filename = os.path.basename(name)
                if not filename: continue
                target_path = os.path.join(output_dir, filename)
                with z.open(name) as source, open(target_path, 'wb') as dest:
                    dest.write(source.read())
                image_map[filename] = target_path
    return image_map

def get_image_rels(doc_part):
    """
    Get relationship map for a document part (e.g. document.part)
    Returns: { rId: target_part }
    """
    return doc_part.rels

def find_images_in_paragraph(paragraph, rels) -> List[Dict]:
    """
    Find images in a paragraph by looking for blip/drawing elements and mapping rId.
    Returns list of dicts: { 'rId': ..., 'filename': ... }
    """
    images = []
    # This is a simplified search. It works for many standard docx images.
    # Namespace for 'r' is usually http://schemas.openxmlformats.org/officeDocument/2006/relationships
    # Look for <a:blip r:embed="rIdX">
    
    # We iterate over the xml to find all blip elements
    blips = paragraph._element.findall('.//{http://schemas.openxmlformats.org/drawingml/2006/main}blip')
    
    # Also check VML for older formats
    vml_imagedata = paragraph._element.findall('.//{urn:schemas-microsoft-com:vml}imagedata')
    
    all_imgs = blips + vml_imagedata
    
    for img_xml in all_imgs:
        embed_attr = img_xml.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
        if not embed_attr:
            embed_attr = img_xml.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id')
        
        if embed_attr and embed_attr in rels:
            rel = rels[embed_attr]
            # Verify it's an image
            if "image" in rel.target_ref:
                # The target_ref is usually like "media/image1.png" relative to the document part
                filename = os.path.basename(rel.target_ref)
                images.append({
                    'rId': embed_attr,
                    'filename': filename
                })
    return images

# --- Structure Extraction ---

def _heading_level_from_style(style) -> int:
    """Infer heading level from style inheritance chain."""
    seen = set()
    current = style
    while current is not None and id(current) not in seen:
        seen.add(id(current))
        name = (getattr(current, "name", "") or "").strip()
        match = re.search(r"heading\s*(\d+)|标题\s*(\d+)", name, re.IGNORECASE)
        if match:
            return int(match.group(1) or match.group(2))
        current = getattr(current, "base_style", None)
    return 0


def _style_font_size_pt(style) -> float | None:
    seen = set()
    current = style
    while current is not None and id(current) not in seen:
        seen.add(id(current))
        try:
            size = current.font.size
        except Exception:
            size = None
        if size is not None:
            try:
                return float(size.pt)
            except Exception:
                pass
        current = getattr(current, "base_style", None)
    return None


def _paragraph_font_size_pt(paragraph) -> float | None:
    sizes: list[float] = []
    for run in getattr(paragraph, "runs", []) or []:
        try:
            size = run.font.size
        except Exception:
            size = None
        if size is not None:
            try:
                sizes.append(float(size.pt))
            except Exception:
                pass
    if sizes:
        return max(sizes)
    return _style_font_size_pt(getattr(paragraph, "style", None))


def _body_font_baseline_pt(doc) -> float | None:
    samples: list[float] = []
    normal_style = None
    try:
        normal_style = doc.styles["Normal"]
    except Exception:
        normal_style = None
    normal_size = _style_font_size_pt(normal_style)
    if normal_size is not None:
        samples.append(normal_size)

    for para in getattr(doc, "paragraphs", []) or []:
        text = (getattr(para, "text", "") or "").strip()
        if not text:
            continue
        if _heading_level_from_style(getattr(para, "style", None)) > 0:
            continue
        size = _paragraph_font_size_pt(para)
        if size is not None:
            samples.append(size)
        if len(samples) >= 200:
            break
    if not samples:
        return None
    samples.sort()
    return samples[len(samples) // 2]


def _is_visual_heading_candidate(paragraph, text: str, body_font_pt: float | None) -> bool:
    if not text or body_font_pt is None:
        return False
    if len(text) > 40:
        return False
    if text.endswith(("；", "。", "，", ";", ".", ",")):
        return False
    para_size = _paragraph_font_size_pt(paragraph)
    if para_size is None:
        return False
    # 通用兜底：字号明显大于正文，视为可单独切片的标题候选。
    return para_size >= body_font_pt + 3.0 or para_size >= body_font_pt * 1.3


def is_heading(paragraph, body_font_pt: float | None = None) -> Tuple[bool, int]:
    """Determine if paragraph is a heading and return level."""
    style_level = _heading_level_from_style(paragraph.style if paragraph else None)
    if style_level > 0:
        return True, style_level
    
    # if hasattr(paragraph, 'paragraph_format') and paragraph.paragraph_format.level is not None:
    #      return True, paragraph.paragraph_format.level
         
    text = paragraph.text.strip()
    
    # 长度和标点校验：避免将长列表项误判为标题
    # 1. 标点符号结尾的通常是列表项，但短标题允许保留
    if text.endswith(('；', '。', '，', ';', '.', ',')):
        if len(text) > 15:
            return False, 0
        
    # 2. Level 5 标题通常较短，放宽长度阈值避免误判
    # 特别针对 （1） 这种格式
    if re.match(r'^（\d+）|^\d+[、.]', text):
        if len(text) > 30:
             return False, 0
    
    # 3. 如果文本包含冒号后跟描述，通常是正文（但短标题允许）
    if '：' in text or ':' in text:
        parts = re.split(r'[:：]', text, 1)
        if len(parts) > 1 and len(parts[1]) > 15: # 冒号后内容较长即视为正文
            return False, 0

    if re.match(r'^第[一二三四五六七八九十\d]+篇', text): return True, 1
    if re.match(r'^第[一二三四五六七八九十\d]+章', text): return True, 2
    if re.match(r'^第[一二三四五六七八九十\d]+节', text): return True, 3
    # 中文大序号（如 一、二、三、）与中文小序号（如 （一）（二））拆成不同层级，
    # 避免（一）覆盖掉上一级“三、...”路径。
    if re.match(r'^[一二三四五六七八九十]+[、.]', text): return True, 4
    if re.match(r'^（[一二三四五六七八九十]+）', text): return True, 5
    if re.match(r'^（\d+）|^\d+[、.]', text): return True, 6
    if _is_visual_heading_candidate(paragraph, text, body_font_pt):
        return True, 4
    
    return False, 0


def _normalize_toc_heading_text(text: str) -> str:
    s = (text or "").replace("\xa0", " ").strip()
    s = re.sub(r"\t+\s*\d+\s*$", "", s)
    s = re.sub(r"\s+\d+\s*$", "", s)
    return s.strip()


def _toc_alias_map(doc) -> Dict[int, Dict[str, str]]:
    alias_map: Dict[int, Dict[str, str]] = {1: {}, 2: {}, 3: {}}
    for para in getattr(doc, "paragraphs", []) or []:
        text = _normalize_toc_heading_text(getattr(para, "text", "") or "")
        if not text:
            continue
        style_name = (getattr(getattr(para, "style", None), "name", "") or "").strip().lower()
        m = re.fullmatch(r"toc\s*([123])", style_name)
        if not m:
            continue
        level = int(m.group(1))
        alias_map[level][text] = text
        prefix_match = re.match(r"^(第[一二三四五六七八九十\d]+[篇章节])\s*(.+)$", text)
        if prefix_match:
            alias_map[level][prefix_match.group(2).strip()] = text
    return alias_map

def process_document(
    docx_path: str,
    api_key: str,
    extract_dir: str = "extracted_images",
    image_model: str = "doubao-seed-1.8",
    image_base_url: str = "https://ark.cn-beijing.volces.com/api/v3",
    image_provider: str = "",
    ark_api_key: str = "",
    volc_ak: str = "",
    volc_sk: str = "",
    ark_project_name: str = "",
    progress_cb=None,
):
    print(f"Reading {docx_path}...")
    doc = Document(docx_path)
    toc_alias_map = _toc_alias_map(doc)
    body_font_pt = _body_font_baseline_pt(doc)
    disable_image_ocr = str(os.getenv("SLICE_DISABLE_IMAGE_OCR", "")).strip().lower() in {
        "1",
        "true",
        "yes",
        "on",
    }
    
    # 1. Extract all images first to a temp dir
    extracted_image_map = extract_images_from_docx(docx_path, extract_dir) # filename -> path
    print(f"Extracted {len(extracted_image_map)} images to {extract_dir}")

    # Prepare logic for processing
    elements = []
    # path_stack 使用 1-based 层级索引，预留到 6 级标题
    path_stack = [None] * 7
    last_mastery = ""
    
    # We iterate over document body elements
    # doc.element.body gives us the order needed
    
    rels = doc.part.rels # Relationship map for the main document part
    
    table_index_global = 0
    image_index_global = 0

    def _build_table_anchor(table_index: int, row_idx: int, col_idx: int, cell_text: str) -> Dict:
        """
        Build a stable anchor describing where an image sits inside a table.
        row_idx/col_idx are 1-based.
        """
        text_preview = (cell_text or "").strip().replace("\n", " ")
        if len(text_preview) > 80:
            text_preview = text_preview[:80] + "..."
        return {
            "anchor_type": "table_cell",
            "table_index": table_index,
            "row": row_idx,
            "col": col_idx,
            "anchor_label": f"表{table_index}-R{row_idx}C{col_idx}",
            "cell_text_preview": text_preview,
        }

    def _render_table_preserve_merge(table) -> Dict[str, object]:
        """
        Render table with merge-awareness.
        - If no merged cells: return Markdown table.
        - If merged cells exist: return HTML table with rowspan/colspan.
        """
        rows = table.rows
        if not rows:
            return {"text": "", "format": "markdown", "has_merged": False}

        # Build grid of underlying XML tc pointers; merged cells share same pointer.
        max_cols = max((len(r.cells) for r in rows), default=0)
        if max_cols == 0:
            return {"text": "", "format": "markdown", "has_merged": False}

        tc_grid = []
        text_grid = []
        for r in rows:
            row_tcs = []
            row_txt = []
            for c in range(max_cols):
                if c < len(r.cells):
                    cell = r.cells[c]
                    row_tcs.append(cell._tc)
                    row_txt.append((cell.text or "").strip().replace("\n", " "))
                else:
                    row_tcs.append(None)
                    row_txt.append("")
            tc_grid.append(row_tcs)
            text_grid.append(row_txt)

        nrows = len(tc_grid)
        ncols = max_cols
        masters = []
        has_merged = False

        for r in range(nrows):
            for c in range(ncols):
                tc = tc_grid[r][c]
                if tc is None:
                    continue
                # Top-left master of a merged block
                is_top = (r == 0) or (tc_grid[r - 1][c] is not tc)
                is_left = (c == 0) or (tc_grid[r][c - 1] is not tc)
                if not (is_top and is_left):
                    continue

                colspan = 1
                while c + colspan < ncols and tc_grid[r][c + colspan] is tc:
                    colspan += 1

                rowspan = 1
                while r + rowspan < nrows and tc_grid[r + rowspan][c] is tc:
                    rowspan += 1

                if rowspan > 1 or colspan > 1:
                    has_merged = True

                masters.append({
                    "r": r,
                    "c": c,
                    "rowspan": rowspan,
                    "colspan": colspan,
                    "text": text_grid[r][c],
                })

        # No merged cells -> output markdown for compactness
        if not has_merged:
            header = [text_grid[0][c] for c in range(ncols)]
            lines = [
                f"| {' | '.join(header)} |",
                f"| {' | '.join(['---'] * ncols)} |",
            ]
            for r in range(1, nrows):
                row_vals = [text_grid[r][c] for c in range(ncols)]
                lines.append(f"| {' | '.join(row_vals)} |")
            return {"text": "\n".join(lines), "format": "markdown", "has_merged": False}

        # Has merged cells -> output HTML with rowspan/colspan
        master_map = {(m["r"], m["c"]): m for m in masters}
        covered = set()
        for m in masters:
            r0, c0 = m["r"], m["c"]
            for rr in range(r0, r0 + m["rowspan"]):
                for cc in range(c0, c0 + m["colspan"]):
                    if rr == r0 and cc == c0:
                        continue
                    covered.add((rr, cc))

        html_lines = ["<table>"]
        for r in range(nrows):
            html_lines.append("  <tr>")
            for c in range(ncols):
                if (r, c) in covered:
                    continue
                m = master_map.get((r, c))
                if not m:
                    continue
                attrs = []
                if m["rowspan"] > 1:
                    attrs.append(f'rowspan="{m["rowspan"]}"')
                if m["colspan"] > 1:
                    attrs.append(f'colspan="{m["colspan"]}"')
                attr_str = (" " + " ".join(attrs)) if attrs else ""
                tag = "th" if r == 0 else "td"
                txt = html.escape(m["text"] or "")
                html_lines.append(f"    <{tag}{attr_str}>{txt}</{tag}>")
            html_lines.append("  </tr>")
        html_lines.append("</table>")
        return {"text": "\n".join(html_lines), "format": "html", "has_merged": True}

    def analyze_extracted_image(fpath: str, fname: str, idx: int, total: int) -> Dict:
        print(f"🖼️ 处理图片 [{idx}/{total}]: {os.path.basename(fpath)}", flush=True)
        analysis = ""
        contains_table = False
        contains_chart = False
        if disable_image_ocr:
            analysis = "(已跳过图片OCR：SLICE_DISABLE_IMAGE_OCR=1)"
        elif api_key and analyze_image_with_qwen_vl:
            analysis_result = analyze_image_with_qwen_vl(
                fpath,
                api_key,
                model_name=image_model,
                base_url=image_base_url,
                provider=image_provider,
                ark_api_key=ark_api_key,
                volc_ak=volc_ak,
                volc_sk=volc_sk,
                ark_project_name=ark_project_name,
            )
            if analysis_result:
                analysis = analysis_result.strip()
                if extract_table_from_content:
                    _, table_content = extract_table_from_content(analysis_result)
                    contains_table = bool(table_content)
                if normalize_image_analysis_content:
                    analysis = normalize_image_analysis_content(analysis, contains_table=contains_table)
                contains_chart = any(
                    kw in analysis_result.lower()
                    for kw in ['坐标', '曲线', '趋势', '图表', '图', 'axis', 'chart']
                )
            else:
                detail = str(getattr(analyze_image_with_qwen_vl, "last_error", "") or "")[:300].strip()
                analysis = (
                    f"(分析失败：{detail})"
                    if detail
                    else "(分析失败：图片模型调用失败，请检查 ARK_API_KEY 或 IMAGE_API_KEY / IMAGE_BASE_URL 配置)"
                )
        elif not api_key:
            analysis = "(分析失败：未配置图片模型 KEY，请在填写您的Key.txt设置 IMAGE_API_KEY 或 OPENAI_API_KEY)"
        else:
            analysis = "(分析失败：图片分析模块未加载)"
        img_obj = {
            "image_id": fname,
            "image_path": fpath,
            "analysis": analysis,
            "contains_table": contains_table,
            "contains_chart": contains_chart,
        }
        if progress_cb:
            progress_cb(
                "image",
                {
                    "index": idx,
                    "total": total,
                    "image_id": fname,
                    "image_path": fpath,
                    "contains_table": contains_table,
                    "contains_chart": contains_chart,
                    "analysis_ok": not str(analysis).startswith("(分析失败"),
                },
            )
        return img_obj
    
    for element in doc.element.body:
        if isinstance(element, CT_P):
            para = Paragraph(element, doc)
            text = para.text.strip()
            
            # Check for images in this paragraph
            # We want to attach images to the content flow
            para_images = find_images_in_paragraph(para, rels)
            
            processed_images = []
            if para_images:
                seen_para_images = set()
                for img_info in para_images:
                    fname = img_info['filename']
                    if fname in seen_para_images:
                        continue
                    seen_para_images.add(fname)
                    fpath = extracted_image_map.get(fname)
                    if fpath:
                        image_index_global += 1
                        total_images = len(extracted_image_map)
                        print(
                            f"🖼️ 处理图片 [{image_index_global}/{total_images}]: {os.path.basename(fpath)}",
                            flush=True,
                        )
                        # Perform analysis if API key is present and analyzer is available
                        analysis = ""
                        contains_table = False
                        contains_chart = False
                        if disable_image_ocr:
                            analysis = "(已跳过图片OCR：SLICE_DISABLE_IMAGE_OCR=1)"
                        elif api_key and analyze_image_with_qwen_vl:
                            analysis_result = analyze_image_with_qwen_vl(
                                fpath,
                                api_key,
                                model_name=image_model,
                                base_url=image_base_url,
                                provider=image_provider,
                                ark_api_key=ark_api_key,
                                volc_ak=volc_ak,
                                volc_sk=volc_sk,
                                ark_project_name=ark_project_name,
                            )
                            if analysis_result:
                                analysis = analysis_result.strip()
                                if extract_table_from_content:
                                    _, table_content = extract_table_from_content(analysis_result)
                                    contains_table = bool(table_content)
                                if normalize_image_analysis_content:
                                    analysis = normalize_image_analysis_content(analysis, contains_table=contains_table)
                                contains_chart = any(
                                    kw in analysis_result.lower()
                                    for kw in ['坐标', '曲线', '趋势', '图表', '图', 'axis', 'chart']
                                )
                            else:
                                detail = ""
                                if analyze_image_with_qwen_vl:
                                    detail = str(getattr(analyze_image_with_qwen_vl, "last_error", "") or "")
                                detail = detail[:300].strip()
                                analysis = (
                                    f"(分析失败：{detail})"
                                    if detail
                                    else "(分析失败：图片模型调用失败，请检查 ARK_API_KEY 或 IMAGE_API_KEY / IMAGE_BASE_URL 配置)"
                                )
                        elif not api_key:
                            analysis = "(分析失败：未配置图片模型 KEY，请在填写您的Key.txt设置 IMAGE_API_KEY 或 OPENAI_API_KEY)"
                        else:
                            analysis = "(分析失败：图片分析模块未加载)"

                        img_obj = {
                           "image_id": fname,
                           "image_path": fpath,
                           "analysis": analysis,
                           "contains_table": contains_table,
                           "contains_chart": contains_chart
                        }
                        if progress_cb:
                            progress_cb(
                                "image",
                                {
                                    "index": image_index_global,
                                    "total": total_images,
                                    "image_id": fname,
                                    "image_path": fpath,
                                    "contains_table": contains_table,
                                    "contains_chart": contains_chart,
                                    "analysis_ok": not str(analysis).startswith("(分析失败"),
                                },
                            )
                        processed_images.append(img_obj)
            
            # Heading Login
            is_head, level = is_heading(para, body_font_pt=body_font_pt)
            if is_head and level > 0:
                # Extract mastery
                m = re.search(r'[（(](掌握|熟悉|了解)[)）]', text)
                if m:
                    last_mastery = m.group(1)
                    text = re.sub(r'[（(](掌握|熟悉|了解)[)）]', '', text).strip()
                text = toc_alias_map.get(level, {}).get(text, text)
                text = _clean_path_seg(text)
                if not text:
                    # Skip empty headings to prevent blank path levels.
                    continue

                path_stack[level] = text
                for i in range(level + 1, 7): path_stack[i] = None
                # Route at most 5 levels (path_stack[1:6]); level-6 headings are content under 5-level
                path_5 = _clean_joined_path(" > ".join([p for p in path_stack[1:6] if p]))
                if level <= 5:
                    elements.append({
                        "type": "heading",
                        "level": level,
                        "text": text,
                        "path": path_5,
                        "mastery": last_mastery
                    })
                else:
                    # Level 6: do not start new slice; emit as paragraph so it merges into current 5-level slice
                    elements.append({
                        "type": "paragraph",
                        "text": text,
                        "path": path_5,
                        "mastery": last_mastery,
                        "images": []
                    })
            else:
                if text or processed_images:
                    current_path = _clean_joined_path(" > ".join([p for p in path_stack[1:6] if p]))
                    elements.append({
                        "type": "paragraph",
                        "text": text,
                        "path": current_path,
                        "mastery": last_mastery,
                        "images": processed_images
                    })

        elif isinstance(element, CT_Tbl):
            # Find the index of this table in the doc.tables list
            # But iterating body elements sequentially matches doc.tables sequentially usually?
            # Actually, `element` IS the table element.
            # We can wrap it in a Table object if needed, or parse XML directly.
            # To be safe, we parse XML or assume sequentiality.
            # doc.tables is a list of all tables.
            # We can try to match by identity? No, `doc.tables` creates new proxy objects.
            # Let's iterate `doc.tables` and pop? No.
            
            import docx.table
            table = docx.table.Table(element, doc)
            table_rendered = _render_table_preserve_merge(table)
            table_text = str(table_rendered.get("text", ""))
            table_format = str(table_rendered.get("format", "markdown"))
            table_has_merged = bool(table_rendered.get("has_merged", False))
            current_path = _clean_joined_path(" > ".join([p for p in path_stack[1:6] if p]))
            table_index_global += 1
            table_images = []
            table_image_anchors = []
            seen_table_images = set()
            table_image_map = {}
            try:
                for row_idx, row in enumerate(table.rows, start=1):
                    for col_idx, cell in enumerate(row.cells, start=1):
                        for para_in_cell in cell.paragraphs:
                            cell_images = find_images_in_paragraph(para_in_cell, rels)
                            for img_info in cell_images:
                                fname = img_info.get("filename")
                                if not fname:
                                    continue
                                fpath = extracted_image_map.get(fname)
                                if not fpath:
                                    continue
                                anchor = _build_table_anchor(table_index_global, row_idx, col_idx, cell.text)
                                table_image_anchors.append({
                                    "image_id": fname,
                                    **anchor,
                                })

                                if fname in table_image_map:
                                    table_image_map[fname].setdefault("anchors", []).append(anchor)
                                    continue

                                if fname not in seen_table_images:
                                    seen_table_images.add(fname)
                                    image_index_global += 1
                                    total_images = len(extracted_image_map)
                                    img_obj = analyze_extracted_image(fpath, fname, image_index_global, total_images)
                                    img_obj["source_type"] = "table_cell"
                                    img_obj["table_index"] = table_index_global
                                    img_obj["anchors"] = [anchor]
                                    table_image_map[fname] = img_obj
            except Exception:
                pass

            table_images = list(table_image_map.values())
            
            elements.append({
                "type": "table",
                "text": table_text,
                "path": current_path,
                "mastery": last_mastery,
                "table_index": table_index_global,
                "table_format": table_format,
                "table_has_merged": table_has_merged,
                "images": table_images,
                "image_anchors": table_image_anchors,
            })

    return elements

def group_and_slice(elements: List[Dict], api_key: str, progress_cb=None):
    # Grouping logic (simplified)
    # Collect all P and Table under the last Heading
    
    slices = []
    current_slice = None
    
    EXCLUDE_KEYWORDS = ["前言", "目录", "相关说明"]
    
    for el in elements:
        if el["type"] == "heading" and 1 <= el["level"] <= 5:  # New slice only for level 1-5; level 6 is content
            # Start new slice
            if current_slice:
                slices.append(current_slice)
            
            current_slice = {
                "完整路径": el["path"],
                "掌握程度": el["mastery"],
                "结构化内容": {
                    "key_params": [],
                    "rules": [],
                    "context_before": "",
                    "tables": [],
                    "context_after": "",
                    "images": [],
                    "image_anchors": [],
                    "formulas": [],
                    "examples": [] # New field for examples
                },
                "metadata": {
                    "类型": "自动组装"
                },
                "_lines_before": [], 
                "_lines_after": [],
                "_hit_visual": False,
                "_in_example": False # Track if we are processing an example block
            }
        
        elif current_slice: # Content
            # Skip if path match excluded
            if any(k in el["path"] for k in EXCLUDE_KEYWORDS):
                continue
            
            if el["type"] == "paragraph":
                text_content = el["text"]
                
                # Check for Example start
                if text_content.startswith("【例】"):
                    current_slice["_in_example"] = True
                    current_slice["结构化内容"]["examples"].append(text_content)
                elif current_slice["_in_example"]:
                    # If we are in an example block, heuristics to decide if we are still in it
                    # Usually example includes question + options + solution (【解】)
                    # For now, simplistic: if it starts with 【解】 or options A/B/C/D or is short/continuation, keep in example
                    # If it looks like a new Heading (handled by outer loop), we exit naturally.
                    current_slice["结构化内容"]["examples"][-1] += "\n" + text_content
                else:
                    # Normal text processing
                    # Check for images FIRST, as they might be inline
                    has_images = False
                    if el.get("images"):
                        current_slice["结构化内容"]["images"].extend(el["images"])
                        current_slice["_hit_visual"] = True
                        has_images = True
                    
                    if text_content:
                        # Check for formulas in text
                        if "=" in text_content and any(x in text_content for x in "+-*/"):
                             current_slice["结构化内容"]["formulas"].append(text_content)
                        
                        # Distribute text
                        if current_slice["_hit_visual"]:
                            current_slice["_lines_after"].append(text_content)
                        else:
                            current_slice["_lines_before"].append(text_content)
            
            elif el["type"] == "table":
                # If we encounter a table, does it belong to the example?
                # If _in_example is True, maybe?
                # For safety, let's reset example mode on visuals for now, or keep separate.
                current_slice["_in_example"] = False 
                current_slice["结构化内容"]["tables"].append(el["text"])
                if el.get("images"):
                    current_slice["结构化内容"]["images"].extend(el["images"])
                if el.get("image_anchors"):
                    current_slice["结构化内容"]["image_anchors"].extend(el["image_anchors"])
                current_slice["_hit_visual"] = True
    
    if current_slice:
        slices.append(current_slice)
        
    # Post-process: Fill context fields
    final_slices = []
    for s in slices:
        s["结构化内容"]["context_before"] = "\n".join(s["_lines_before"])
        s["结构化内容"]["context_after"] = "\n".join(s["_lines_after"])
        # Cleanup temp fields
        del s["_lines_before"]
        del s["_lines_after"]
        del s["_hit_visual"]
        s.pop("_in_example", None)
        
        # Metadata updates
        s["metadata"]["表格索引"] = len(s["结构化内容"]["tables"])
        s["metadata"]["图片索引"] = len(s["结构化内容"]["images"])
        s["metadata"]["表格图片锚点数"] = len(s["结构化内容"].get("image_anchors", []) or [])
        s["metadata"]["包含计算公式"] = len(s["结构化内容"]["formulas"]) > 0
        if progress_cb:
            progress_cb(
                "slice_draft",
                {
                    "path": s.get("完整路径", ""),
                    "images": len(s["结构化内容"].get("images", []) or []),
                    "tables": len(s["结构化内容"].get("tables", []) or []),
                    "examples": len(s["结构化内容"].get("examples", []) or []),
                },
            )
        
        final_slices.append(s)
        
    return final_slices


_PATH_L4_RE = re.compile(r"^(?:[一二三四五六七八九十百千万]+[、.]|\d+[、.])")
_PATH_CHILD_RE = re.compile(r"^(?:（[一二三四五六七八九十百千万]+）|（\d+）)")
_INVISIBLE_RE = re.compile(r"[\u200b\u200c\u200d\ufeff\u2060]")


def _clean_path_seg(seg: str) -> str:
    s = str(seg or "")
    s = _INVISIBLE_RE.sub("", s)
    return s.strip()


def _clean_joined_path(path: str) -> str:
    segs = [_clean_path_seg(x) for x in str(path or "").split(" > ")]
    segs = [x for x in segs if x]
    return " > ".join(segs)


# Minimum content length (chars) for a valid slice; below this we merge at parent level
MIN_SLICE_CONTENT_CHARS = 200
MAX_MERGED_SLICE_CONTENT_CHARS = 2500


def _slice_content_length(s: Dict) -> int:
    """Total character count of slice content (context, tables, examples, formulas)."""
    content = s.get("结构化内容", {}) or {}
    parts = [
        content.get("context_before", "") or "",
        content.get("context_after", "") or "",
        "\n".join(content.get("tables", []) or []),
        "\n".join(content.get("examples", []) or []),
        " ".join(content.get("formulas", []) or []),
    ]
    return sum(len(str(p)) for p in parts)


def _parent_path(path: str) -> Optional[str]:
    """Return path with last segment removed (上一级), or None if single segment."""
    p = _clean_joined_path(path or "")
    segs = [x for x in p.split(" > ") if x.strip()]
    if len(segs) <= 1:
        return None
    return " > ".join(segs[:-1])


def _first_child_under_parent(path: str, parent: str) -> str:
    """Return immediate child segment under parent path, or '' when unavailable."""
    p = _clean_joined_path(path or "")
    base = _clean_joined_path(parent or "")
    if not p or not base:
        return ""
    prefix = base + " > "
    if not p.startswith(prefix):
        return ""
    rel = p[len(prefix):].strip()
    if not rel:
        return ""
    return rel.split(" > ")[0].strip()


def merge_short_slices_at_parent(
    slices: List[Dict],
    min_chars: int = MIN_SLICE_CONTENT_CHARS,
    max_merged_chars: int = MAX_MERGED_SLICE_CONTENT_CHARS,
) -> List[Dict]:
    """
    After slicing: if any slice has content < min_chars, invalidate all slices under
    that path's parent (上一级), and replace them with a single re-sliced slice at the
    parent level. The new slice uses the parent path and merged content so each
    slice is at least min_chars (by merging siblings at parent dimension).
    """
    if not slices:
        return slices
    # Normalize path for each slice
    indexed = []
    for i, s in enumerate(slices):
        if not isinstance(s, dict):
            indexed.append((i, s, None))
            continue
        path = _clean_joined_path(str(s.get("完整路径", "") or ""))
        length = _slice_content_length(s)
        indexed.append((i, s, {"path": path, "length": length}))
    # Merge branches that contain short slices:
    # key = (parent_path, child_seg). This avoids swallowing all siblings under a parent.
    branches_to_merge: set[tuple[str, str]] = set()
    for _, s, meta in indexed:
        if meta and meta["length"] < min_chars:
            parent = _parent_path(meta["path"])
            if parent:
                child = _first_child_under_parent(meta["path"], parent)
                branches_to_merge.add((parent, child))
    if not branches_to_merge:
        return slices

    # For each selected branch, collect slice indices under that branch only.
    to_skip: set = set()
    merged_by_first: Dict[int, List[Dict]] = {}
    for parent_path, child_seg in sorted(branches_to_merge):
        target_path = parent_path if not child_seg else _clean_joined_path(f"{parent_path} > {child_seg}")
        group_indices = []
        group_slices = []
        for i, s, meta in indexed:
            if not meta:
                continue
            path = meta["path"]
            if path == target_path or path.startswith(target_path + " > "):
                group_indices.append(i)
                group_slices.append(s)
        if not group_slices:
            continue
        to_skip.update(group_indices)
        first_idx = min(group_indices)
        merged = _merge_slices_at_path(group_slices, target_path, max_merged_chars=max_merged_chars)
        merged_by_first[first_idx] = merged
    # Build output: skip merged members, emit merged slice at first index of each group
    out = []
    for i, s, _ in indexed:
        if i in to_skip:
            if i in merged_by_first:
                out.extend(merged_by_first[i])
            continue
        out.append(s)
    return out


def _merge_slices_at_path(slice_list: List[Dict], path: str, max_merged_chars: int = MAX_MERGED_SLICE_CONTENT_CHARS) -> List[Dict]:
    """Merge slices at parent level and re-group into 1..N slices (avoid one oversized slice)."""
    if not slice_list:
        raise ValueError("slice_list must be non-empty")
    first = slice_list[0]

    def _relative_route(src_path: str, base_path: str) -> str:
        sp = _clean_joined_path(src_path or "")
        bp = _clean_joined_path(base_path or "")
        if not sp:
            return ""
        if not bp or sp == bp:
            return ""
        prefix = bp + " > "
        if sp.startswith(prefix):
            return sp[len(prefix):].strip()
        return sp

    # Convert each source slice into an ordered unit to preserve semantics, then re-pack by size.
    units: List[Dict] = []
    for s in slice_list:
        c = s.get("结构化内容", {}) or {}
        src_path = str(s.get("完整路径", "") or "")
        route_hint = _relative_route(src_path, path)
        before = str(c.get("context_before", "") or "").strip()
        after = str(c.get("context_after", "") or "").strip()
        tables = list(c.get("tables", []) or [])
        images = list(c.get("images", []) or [])
        anchors = list(c.get("image_anchors", []) or [])
        formulas = list(c.get("formulas", []) or [])
        examples = list(c.get("examples", []) or [])
        preview_len = len("\n".join([before, "\n".join(tables), after, "\n".join(examples), " ".join(formulas)]))
        units.append(
            {
                "route_hint": route_hint,
                "before": before,
                "after": after,
                "tables": tables,
                "images": images,
                "anchors": anchors,
                "formulas": formulas,
                "examples": examples,
                "preview_len": preview_len,
            }
        )

    # Pack units into buckets to avoid one oversized merged slice.
    buckets: List[List[Dict]] = []
    cur: List[Dict] = []
    cur_len = 0
    for u in units:
        u_len = int(u.get("preview_len", 0) or 0)
        if cur and cur_len + u_len > max_merged_chars:
            buckets.append(cur)
            cur = [u]
            cur_len = u_len
        else:
            cur.append(u)
            cur_len += u_len
    if cur:
        buckets.append(cur)
    if not buckets:
        buckets = [units]

    out: List[Dict] = []
    mastery = first.get("掌握程度", "") or "了解"
    total = len(buckets)
    for idx, bucket in enumerate(buckets, start=1):
        content = {
            "key_params": [],
            "rules": [],
            "context_before": "",
            "tables": [],
            "context_after": "",
            "images": [],
            "image_anchors": [],
            "formulas": [],
            "examples": [],
        }
        before_parts = []
        after_parts = []
        for u in bucket:
            route_hint = str(u.get("route_hint", "") or "").strip()
            if route_hint:
                before_parts.append(f"【路由】{route_hint}")
            if u.get("before"):
                before_parts.append(str(u["before"]).strip())
            for t in (u.get("tables") or []):
                content["tables"].append(t)
            for img in (u.get("images") or []):
                content["images"].append(img)
            for anc in (u.get("anchors") or []):
                content["image_anchors"].append(anc)
            for f in (u.get("formulas") or []):
                content["formulas"].append(f)
            for ex in (u.get("examples") or []):
                content["examples"].append(ex)
            if u.get("after"):
                after_parts.append(str(u["after"]).strip())
        content["context_before"] = "\n\n".join(before_parts).strip()
        content["context_after"] = "\n\n".join(after_parts).strip()

        meta = dict(first.get("metadata") or {})
        meta["类型"] = "自动组装"
        meta["merged_from_short_slices"] = True
        meta["merged_count"] = len(slice_list)
        meta["reclustered_after_merge"] = total > 1
        meta["recluster_index"] = idx
        meta["recluster_total"] = total
        meta["表格索引"] = len(content["tables"])
        meta["图片索引"] = len(content["images"])
        meta["表格图片锚点数"] = len(content.get("image_anchors", []) or [])
        meta["包含计算公式"] = len(content["formulas"]) > 0

        # Keep textbook-authored route only; do not emit synthetic route segments.
        slice_path = path
        out.append(
            {
                "完整路径": slice_path,
                "掌握程度": mastery,
                "结构化内容": content,
                "metadata": meta,
            }
        )
    return out


def repair_flattened_paths(slices: List[Dict]) -> List[Dict]:
    """
    Repair flattened paths like:
    ... > （一）周期影响因素
    to:
    ... > 三、房地产市场周期波动 > （一）周期影响因素

    This runs as a source-side safeguard before writing JSONL.
    """
    latest_l4_by_p3: Dict[str, str] = {}
    out: List[Dict] = []
    for s in slices:
        if not isinstance(s, dict):
            out.append(s)
            continue
        path = _clean_joined_path(str(s.get("完整路径", "") or ""))
        segs = [_clean_path_seg(x) for x in path.split(" > ") if _clean_path_seg(x)]
        if len(segs) >= 4:
            p3 = " > ".join(segs[:3])
            seg4 = segs[3]
            if _PATH_L4_RE.match(seg4):
                latest_l4_by_p3[p3] = seg4
            elif _PATH_CHILD_RE.match(seg4):
                parent = latest_l4_by_p3.get(p3)
                if parent:
                    segs = [*segs[:3], parent, *segs[3:]]
        fixed_path = _clean_joined_path(" > ".join(segs) if segs else path)
        if fixed_path != path:
            patched = dict(s)
            patched["完整路径"] = fixed_path
            out.append(patched)
        else:
            out.append(s)
    return out


# --- Optional: Embedding-based sub-slicing under level-5 paths ---

_BGE_MODEL = None
_BGE_MODEL_NAME = "BAAI/bge-small-zh-v1.5"


def _get_bge_model():
    """Lazy-load BGE model for sub-slicing. Returns None if unavailable."""
    global _BGE_MODEL
    if _BGE_MODEL is not None:
        return _BGE_MODEL
    if SentenceTransformer is None:
        return None
    try:
        _BGE_MODEL = SentenceTransformer(_BGE_MODEL_NAME)
        return _BGE_MODEL
    except Exception:
        return None


def _path_depth_5(path: str) -> bool:
    segs = [x.strip() for x in str(path or "").split(" > ") if x.strip()]
    return len(segs) == 5


def _slice_blocks_for_embedding(s: Dict) -> List[Dict[str, str]]:
    """
    Convert a slice's structured content into ordered blocks for embedding.
    Each block: {"type": "text"|"table"|"example", "text": "..."}.
    """
    content = s.get("结构化内容", {}) or {}
    blocks: List[Dict[str, str]] = []

    def _add_text(txt: str):
        t = str(txt or "").strip()
        if not t:
            return
        # Split by blank lines to keep local cohesion.
        parts = [p.strip() for p in re.split(r"\n\s*\n+", t) if p.strip()]
        for p in parts:
            blocks.append({"type": "text", "text": p})

    _add_text(content.get("context_before", ""))
    for ex in (content.get("examples", []) or []):
        t = str(ex or "").strip()
        if t:
            blocks.append({"type": "example", "text": t})
    for tb in (content.get("tables", []) or []):
        t = str(tb or "").strip()
        if t:
            blocks.append({"type": "table", "text": t})
    _add_text(content.get("context_after", ""))
    return blocks


def _encode_blocks(blocks: List[Dict[str, str]]) -> Optional["np.ndarray"]:
    """Encode blocks with BGE. Returns (N, dim) normalized embeddings or None."""
    if not blocks:
        return None
    if np is None:
        return None
    model = _get_bge_model()
    if model is None:
        return None
    texts = [b.get("text", "") for b in blocks]
    embs = model.encode(texts, batch_size=64, normalize_embeddings=True)
    return np.asarray(embs, dtype=np.float32)


def _build_subslice_from_blocks(base: Dict, blocks: List[Dict[str, str]], sub_index: int, total: int) -> Dict:
    """Build a new slice dict from a subset of ordered blocks."""
    base_path = str(base.get("完整路径", "") or "")
    mastery = base.get("掌握程度", "")
    new_slice = {
        "完整路径": base_path,
        "掌握程度": mastery,
        "结构化内容": {
            "key_params": [],
            "rules": [],
            "context_before": "",
            "tables": [],
            "context_after": "",
            "images": [],
            "image_anchors": [],
            "formulas": [],
            "examples": [],
        },
        "metadata": dict((base.get("metadata") or {})),
    }
    # Keep visuals/formulas on all subslices for safety (do not drop information).
    base_content = base.get("结构化内容", {}) or {}
    new_slice["结构化内容"]["images"] = list(base_content.get("images", []) or [])
    new_slice["结构化内容"]["image_anchors"] = list(base_content.get("image_anchors", []) or [])
    new_slice["结构化内容"]["formulas"] = list(base_content.get("formulas", []) or [])

    # Refill content from blocks while preserving order roughly.
    text_parts: List[str] = []
    after_parts: List[str] = []
    hit_visual = False
    for b in blocks:
        btype = b.get("type")
        btext = str(b.get("text", "") or "").strip()
        if not btext:
            continue
        if btype == "table":
            new_slice["结构化内容"]["tables"].append(btext)
            hit_visual = True
        elif btype == "example":
            new_slice["结构化内容"]["examples"].append(btext)
        else:
            if hit_visual:
                after_parts.append(btext)
            else:
                text_parts.append(btext)
    new_slice["结构化内容"]["context_before"] = "\n\n".join(text_parts).strip()
    new_slice["结构化内容"]["context_after"] = "\n\n".join(after_parts).strip()

    # Metadata for traceability
    new_slice["metadata"]["subslice_enabled"] = True
    new_slice["metadata"]["subslice_index"] = sub_index
    new_slice["metadata"]["subslice_total"] = total
    # Provide a lightweight title hint (first non-empty line)
    preview_src = (
        (new_slice["结构化内容"]["context_before"] or "").splitlines()
        or (new_slice["结构化内容"]["context_after"] or "").splitlines()
    )
    preview = (preview_src[0].strip() if preview_src else "")
    if preview:
        if len(preview) > 40:
            preview = preview[:40] + "..."
        new_slice["metadata"]["subslice_title_hint"] = preview

    # Update metadata indices
    new_slice["metadata"]["表格索引"] = len(new_slice["结构化内容"]["tables"])
    new_slice["metadata"]["图片索引"] = len(new_slice["结构化内容"]["images"])
    new_slice["metadata"]["表格图片锚点数"] = len(new_slice["结构化内容"].get("image_anchors", []) or [])
    new_slice["metadata"]["包含计算公式"] = len(new_slice["结构化内容"]["formulas"]) > 0
    return new_slice


def apply_embedding_subslicing(
    slices: List[Dict],
    enabled: bool = True,
    split_sim_threshold: float = 0.75,
    merge_sim_threshold: float = 0.85,
    short_total_chars_threshold: int = 800,
) -> List[Dict]:
    """
    For slices whose path depth is exactly 5, optionally sub-slice content using BGE embeddings.
    - Split: create a new subslice when consecutive block similarity < split_sim_threshold.
    - Merge: if total chars < short_total_chars_threshold and average similarity >= merge_sim_threshold,
      keep as a single slice (no sub-slicing).
    If embedding deps are unavailable, returns original slices.
    """
    if not enabled:
        return slices
    if np is None or SentenceTransformer is None:
        return slices
    if _get_bge_model() is None:
        return slices

    out: List[Dict] = []
    for s in slices:
        if not isinstance(s, dict):
            out.append(s)
            continue
        if not _path_depth_5(s.get("完整路径", "")):
            out.append(s)
            continue

        blocks = _slice_blocks_for_embedding(s)
        if len(blocks) <= 1:
            out.append(s)
            continue

        total_chars = sum(len(b.get("text", "") or "") for b in blocks)
        embs = _encode_blocks(blocks)
        if embs is None or embs.shape[0] != len(blocks):
            out.append(s)
            continue

        sims = (embs[:-1] * embs[1:]).sum(axis=1)  # cosine similarity (normalized embeddings)
        avg_sim = float(sims.mean()) if sims.size else 1.0
        if total_chars < short_total_chars_threshold and avg_sim >= merge_sim_threshold:
            # Short and highly coherent -> keep as single slice.
            s2 = dict(s)
            s2.setdefault("metadata", {})
            s2["metadata"]["subslice_enabled"] = True
            s2["metadata"]["subslice_merged"] = True
            s2["metadata"]["subslice_reason"] = "short_and_coherent"
            out.append(s2)
            continue

        # Split by boundary sims
        groups: List[List[Dict[str, str]]] = []
        cur: List[Dict[str, str]] = [blocks[0]]
        for i in range(1, len(blocks)):
            if float(sims[i - 1]) < split_sim_threshold:
                groups.append(cur)
                cur = [blocks[i]]
            else:
                cur.append(blocks[i])
        if cur:
            groups.append(cur)

        if len(groups) <= 1:
            s2 = dict(s)
            s2.setdefault("metadata", {})
            s2["metadata"]["subslice_enabled"] = True
            s2["metadata"]["subslice_merged"] = True
            s2["metadata"]["subslice_reason"] = "no_split_boundary"
            out.append(s2)
            continue

        for idx, g in enumerate(groups, start=1):
            out.append(_build_subslice_from_blocks(s, g, sub_index=idx, total=len(groups)))
    return out

def _parse_formula_table(table_text: str) -> List[str]:
    """Parse a Markdown table for formulas (编号/公式). Returns list of formula strings."""
    formulas = []
    lines = [ln.strip() for ln in (table_text or "").splitlines() if ln.strip()]
    if not lines:
        return formulas
    # Expect header row containing 编号 and 公式
    header = lines[0]
    if "编号" not in header or "公式" not in header:
        return formulas
    # Skip header + separator
    for ln in lines[2:]:
        if "|" not in ln:
            continue
        parts = [p.strip() for p in ln.strip("|").split("|")]
        if len(parts) < 2:
            continue
        formula = parts[1]
        if formula:
            formulas.append(formula)
    return formulas

def _assign_formulas_to_slices(final_slices: List[Dict], appendix_slice: Dict, formulas: List[str]) -> None:
    """Assign formulas to best-matching slices based on left-side keyword."""
    if not formulas or not final_slices:
        return
    # Build searchable content for each slice
    slice_index = []
    for s in final_slices:
        content = s.get("结构化内容", {})
        text = " ".join([
            s.get("完整路径", ""),
            content.get("context_before", ""),
            content.get("context_after", ""),
            "\n".join(content.get("tables", []) or []),
            "\n".join(content.get("examples", []) or []),
        ])
        slice_index.append((s, text))

    appendix_text = ""
    if appendix_slice:
        appendix_text = " ".join([
            appendix_slice.get("完整路径", ""),
            appendix_slice.get("结构化内容", {}).get("context_before", ""),
            appendix_slice.get("结构化内容", {}).get("context_after", ""),
        ])

    for fml in formulas:
        left = fml.split("=", 1)[0].strip()
        if not left:
            continue
        best = None
        best_score = 0
        for s, text in slice_index:
            score = 0
            if left in text:
                score += 2
            if "附录" in text:
                score += 1
            if score > best_score:
                best_score = score
                best = s
        target = best if best_score > 0 else appendix_slice
        if not target:
            continue
        target.setdefault("结构化内容", {}).setdefault("formulas", [])
        if fml not in target["结构化内容"]["formulas"]:
            target["结构化内容"]["formulas"].append(fml)

def main():
    parser = argparse.ArgumentParser(description="Generate knowledge slices from textbook docx")
    parser.add_argument("--tenant-id", default="", help="城市租户ID，例如 hz/bj/sh")
    parser.add_argument("--docx", default="第26届存量房教材模板-勘误版0912-干净版.docx")
    parser.add_argument("--output", default="")
    parser.add_argument("--extract-dir", default="extracted_images", help="图片提取目录")
    args = parser.parse_args()

    config = load_config()
    # 图片分析走 provider 分流：
    # - provider=ait: 优先 AIT_API_KEY / AIT_BASE_URL
    # - provider=ark: 优先 Ark 链路配置
    image_model = config.get("IMAGE_MODEL") or "doubao-seed-1.8"
    image_provider = (config.get("IMAGE_PROVIDER") or "").lower()
    if image_provider == "ark":
        api_key = (
            config.get("IMAGE_API_KEY")
            or config.get("ARK_API_KEY")
            or config.get("OPENAI_API_KEY")
            or ""
        )
        image_base_url = (
            config.get("IMAGE_BASE_URL")
            or config.get("ARK_BASE_URL")
            or "https://ark.cn-beijing.volces.com/api/v3"
        )
    else:
        api_key = (
            config.get("AIT_API_KEY")
            or config.get("IMAGE_API_KEY")
            or config.get("OPENAI_API_KEY")
            or config.get("CRITIC_API_KEY")
            or ""
        )
        image_base_url = (
            config.get("IMAGE_BASE_URL")
            or config.get("AIT_BASE_URL")
            or config.get("OPENAI_BASE_URL")
            or "https://openapi-ait.ke.com/v1"
        )
    ark_api_key = config.get("ARK_API_KEY") or ""
    volc_ak = config.get("VOLC_ACCESS_KEY_ID") or ""
    volc_sk = config.get("VOLC_SECRET_ACCESS_KEY") or ""
    ark_project_name = config.get("ARK_PROJECT_NAME") or ""
    
    docx_path = args.docx
    if args.output:
        output_file = args.output
    elif args.tenant_id:
        output_file = str(tenant_slices_dir(args.tenant_id) / "knowledge_slices.jsonl")
    else:
        output_file = "test_knowledge_slices.jsonl"
    progress_file = f"{output_file}.progress.jsonl"
    os.makedirs(os.path.dirname(progress_file) or ".", exist_ok=True)
    with open(progress_file, "w", encoding="utf-8"):
        pass

    def _emit_progress(event_type: str, payload: Dict):
        row = {
            "ts": datetime.now(timezone.utc).isoformat(),
            "event": event_type,
            "payload": payload or {},
        }
        with open(progress_file, "a", encoding="utf-8") as pf:
            pf.write(json.dumps(row, ensure_ascii=False) + "\n")
            pf.flush()
    
    print("🚀 Starting Logic Extraction...")
    _emit_progress("start", {"docx": docx_path, "output_file": output_file})
    elements = process_document(
        docx_path,
        api_key,
        args.extract_dir,
        image_model=image_model,
        image_base_url=image_base_url,
        image_provider=image_provider,
        ark_api_key=ark_api_key,
        volc_ak=volc_ak,
        volc_sk=volc_sk,
        ark_project_name=ark_project_name,
        progress_cb=_emit_progress,
    )
    print(f"Stats: {len(elements)} elements found.")
    _emit_progress("elements_ready", {"count": len(elements)})
    
    slices = group_and_slice(elements, api_key, progress_cb=_emit_progress)
    _emit_progress("slices_grouped", {"count": len(slices)})

    # Appendix formula table split + formula assignment
    appendix_slice = None
    appendix_formulas = []
    appendix_table = None
    source_slice = None
    for s in slices:
        content = s.get("结构化内容", {})
        tables = content.get("tables", []) or []
        for t in tables:
            if "计算公式汇总表" in t and "编号" in t and "公式" in t:
                appendix_formulas = _parse_formula_table(t)
                appendix_table = t
                source_slice = s
                break
        if appendix_table:
            break

    if appendix_table:
        # Remove table from source slice
        source_tables = source_slice.get("结构化内容", {}).get("tables", [])
        source_slice["结构化内容"]["tables"] = [t for t in source_tables if t != appendix_table]
        # Remove appendix marker line from context_after if present
        ca = source_slice.get("结构化内容", {}).get("context_after", "")
        if "附录" in ca and "计算公式汇总表" in ca:
            source_slice["结构化内容"]["context_after"] = ca.replace("附录  计算公式汇总表", "").replace("附录 计算公式汇总表", "").strip()

        appendix_slice = {
            "完整路径": "附录  计算公式汇总表",
            "掌握程度": "了解",
            "结构化内容": {
                "key_params": [],
                "rules": [],
                "context_before": "",
                "tables": [appendix_table],
                "context_after": "",
                "images": [],
                "image_anchors": [],
                "formulas": [],
                "examples": []
            },
            "metadata": {
                "类型": "自动组装"
            },
            "_in_example": False
        }
        slices.append(appendix_slice)
        _assign_formulas_to_slices(slices, appendix_slice, appendix_formulas)

    # Source-side guard: do not output flattened child headings.
    slices = repair_flattened_paths(slices)

    # Optional: embedding-based sub-slicing under level-5 paths.
    # This is best-effort: if embedding deps/model are unavailable, it is skipped silently.
    # Thresholds can be tuned later; keep conservative defaults to avoid over-splitting.
    slices = apply_embedding_subslicing(
        slices,
        enabled=True,
        split_sim_threshold=float(config.get("SLICE_EMB_SPLIT_SIM", "") or 0.75),
        merge_sim_threshold=float(config.get("SLICE_EMB_MERGE_SIM", "") or 0.85),
        short_total_chars_threshold=int(config.get("SLICE_EMB_SHORT_CHARS", "") or 800),
    )

    # Keep original slicing output; short-slice fallback merge is disabled by product decision.

    # Filter out empty slices (often TOC or empty headers)
    valid_slices = []
    for s in slices:
        content = s["结构化内容"]
        has_content = (
            content["context_before"].strip() or 
            content["context_after"].strip() or 
            content["tables"] or 
            content["images"] or 
            content["formulas"] or
            content["examples"]
        )
        if has_content and not is_toc_slice(s):
            valid_slices.append(s)
            
    print(f"Generated {len(valid_slices)} slices (filtered {len(slices) - len(valid_slices)} empty/TOC slices). Saving to {output_file}...")
    _emit_progress(
        "slices_filtered",
        {
            "valid_count": len(valid_slices),
            "filtered_count": len(slices) - len(valid_slices),
        },
    )
    
    with open(output_file, 'w', encoding='utf-8') as f:
        total_valid = len(valid_slices)
        for idx, s in enumerate(valid_slices, 1):
            f.write(json.dumps(s, ensure_ascii=False) + "\n")
            _emit_progress(
                "slice_final_written",
                {
                    "index": idx,
                    "total": total_valid,
                    "path": s.get("完整路径", ""),
                    "images": len((s.get("结构化内容", {}) or {}).get("images", []) or []),
                },
            )
    _emit_progress("done", {"output_file": output_file, "count": len(valid_slices)})
            
    print("✅ Done.")

def is_toc_line(line: str) -> bool:
    stripped = line.strip()
    if not stripped:
        return False
    if "目录" in stripped or "目 录" in stripped:
        return True
    if "\t" in line and re.search(r"\d+\s*$", stripped):
        return True
    if re.search(r"(·{2,}|\.{2,}|…{2,})", stripped) and re.search(r"\d+\s*$", stripped):
        return True
    if re.search(r"第.+(章|节|篇|部分|单元).*\d+\s*$", stripped) and len(stripped) <= 40:
        return True
    return False

def is_toc_slice(slice_obj: dict) -> bool:
    content = slice_obj.get("结构化内容", {})
    if content.get("tables") or content.get("images") or content.get("formulas"):
        return False
    path = slice_obj.get("完整路径", "") or ""
    if "目录" in path or "目 录" in path:
        return True
    text = "\n".join([
        content.get("context_before", ""),
        content.get("context_after", ""),
    ])
    lines = [ln for ln in text.splitlines() if ln.strip()]
    if not lines:
        return False
    if any("目录" in ln or "目 录" in ln for ln in lines):
        return True
    return all(is_toc_line(ln) for ln in lines)

if __name__ == '__main__':
    main()
