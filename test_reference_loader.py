import json

from docx import Document

from exam_factory import KnowledgeRetriever
from reference_loader import load_reference_questions


def test_load_reference_questions_from_docx(tmp_path):
    doc_path = tmp_path / "reference.docx"
    doc = Document()
    doc.add_paragraph("1. 关于存量房交易流程，下列说法正确的是")
    doc.add_paragraph("A. 先网签")
    doc.add_paragraph("B. 先缴税")
    doc.add_paragraph("答案：A")
    doc.add_paragraph("解析：教材明确要求先网签。")
    doc.save(str(doc_path))

    df = load_reference_questions(doc_path)

    assert len(df) == 1
    assert df.iloc[0]["题干"] == "关于存量房交易流程，下列说法正确的是"
    assert df.iloc[0]["选项1"] == "先网签"
    assert df.iloc[0]["正确答案"] == "A"
    assert "先网签" in df.iloc[0]["解析"]


def test_knowledge_retriever_allows_missing_reference_file(tmp_path):
    kb_path = tmp_path / "kb.jsonl"
    kb_path.write_text(
        json.dumps(
            {
                "完整路径": "第1篇 > 第1章 > 第1节 > 知识点A",
                "核心内容": "教材内容A",
                "结构化内容": {"examples": []},
            },
            ensure_ascii=False,
        )
        + "\n",
        encoding="utf-8",
    )

    missing_history = tmp_path / "missing.docx"
    retriever = KnowledgeRetriever(str(kb_path), str(missing_history), mapping_path=str(tmp_path / "missing_mapping.json"))

    assert retriever.history_df.empty
    assert retriever.get_examples_by_knowledge_point({"完整路径": "第1篇 > 第1章 > 第1节 > 知识点A"}) == []
