from __future__ import annotations

import argparse
import json
import math
import re
from collections import Counter
from datetime import datetime
from pathlib import Path
from typing import Any

import pandas as pd
from docx import Document
from docx.shared import Pt
from dotenv import load_dotenv
from langchain_core.runnables import RunnableLambda

from src.llm import ReliableLLMClient, build_llm
from src.pipeline.graph import run_judge
from src.prompt_loader import load_prompt_pair
from src.schemas.evaluation import QuestionInput


def _compute_batch_rule_checks(questions: list[QuestionInput]) -> list[dict[str, str]]:
    """套卷级规则判定（阈值告警，不只是统计展示）。"""
    checks: list[dict[str, str]] = []

    # 规则1：判断题正误比例（接近1:1）
    tf_items = [q for q in questions if q.question_type == "true_false"]
    def _tf_answer_to_bool(q: QuestionInput) -> str:
        ans = str(q.correct_answer or "").strip().upper()
        if ans in {"正确", "错误"}:
            return ans
        if ans in {"A", "B"} and len(q.options) >= 2:
            a_opt = _to_str(q.options[0])
            b_opt = _to_str(q.options[1])
            if "正确" in a_opt and "错误" in b_opt:
                return "正确" if ans == "A" else "错误"
            if "错误" in a_opt and "正确" in b_opt:
                return "错误" if ans == "A" else "正确"
        return ""

    tf_true = sum(1 for q in tf_items if _tf_answer_to_bool(q) == "正确")
    tf_false = sum(1 for q in tf_items if _tf_answer_to_bool(q) == "错误")
    if tf_items:
        total_tf = len(tf_items)
        # 样本较小时放宽，样本>=6时按0.3~0.7判定
        lower, upper = (0.2, 0.8) if total_tf < 6 else (0.3, 0.7)
        true_ratio = tf_true / total_tf if total_tf else 0.0
        passed = lower <= true_ratio <= upper
        checks.append(
            {
                "name": "判断题正误比例",
                "status": "PASS" if passed else "ALERT",
                "detail": f"正确={tf_true}，错误={tf_false}，正确占比={true_ratio:.2f}（阈值 {lower:.1f}~{upper:.1f}）",
            }
        )

    # 规则2：多选ABCD全对上限（<=5）
    multi_items = [q for q in questions if q.question_type == "multiple_choice"]
    full_abcd = 0
    for q in multi_items:
        ans = str(q.correct_answer or "").upper().replace("，", ",").replace("、", ",")
        toks = sorted(set([x.strip() for x in ans.split(",") if x.strip()]))
        if toks == ["A", "B", "C", "D"]:
            full_abcd += 1
    if multi_items:
        passed = full_abcd <= 5
        checks.append(
            {
                "name": "多选ABCD全对数量上限",
                "status": "PASS" if passed else "ALERT",
                "detail": f"ABCD全对数量={full_abcd}（阈值 <= 5）",
            }
        )

    # 规则3：答案字母分布均衡（A/B/C/D）
    answer_letters = Counter()
    for q in questions:
        ans = str(q.correct_answer or "").upper().replace("，", ",").replace("、", ",")
        for t in [x.strip() for x in ans.split(",") if x.strip() in {"A", "B", "C", "D"}]:
            answer_letters[t] += 1
    if answer_letters:
        total = sum(answer_letters.values())
        max_cnt = max(answer_letters.values())
        min_cnt = min(answer_letters.values())
        # 宽松判定：最大最小占比差 <= 0.35
        diff_ratio = (max_cnt - min_cnt) / total if total else 0.0
        passed = diff_ratio <= 0.35
        checks.append(
            {
                "name": "答案字母分布均衡",
                "status": "PASS" if passed else "ALERT",
                "detail": f"分布={dict(answer_letters)}，最大最小差占比={diff_ratio:.2f}（阈值 <= 0.35）",
            }
        )

    return checks


def _compute_batch_llm_negation_check(
    questions: list[QuestionInput],
    llm: Any,
) -> dict[str, str] | None:
    """组卷级 LLM 语义判定：否定设问占比应偏低。

    注意：这里刻意要求模型“按语义判断”，不按“不/非”等关键词机械匹配。
    """
    if not questions or llm is None:
        return None

    items = "\n".join([f"{q.question_id}: {q.stem}" for q in questions])
    default_system = "你是房地产考试命题审核员。请判断整套题中“否定设问”的占比是否过高。"
    default_human = "题目列表：\n{items}\n\n请输出JSON。"
    system_prompt, human_prompt = load_prompt_pair(
        "prompts/layer4_batch_negation_check.md",
        default_system,
        default_human,
        ["items"],
    )
    prompt = f"{system_prompt}\n\n{human_prompt.format(items=items)}"

    client = ReliableLLMClient(llm, timeout_seconds=120, retries=1)
    data = client.invoke_json(
        prompt,
        fallback={"negation_like_ids": [], "semantic_basis": ""},
    )

    ids_raw = data.get("negation_like_ids") or []
    valid_ids = {q.question_id for q in questions}
    neg_ids = [str(x).strip() for x in ids_raw if str(x).strip() in valid_ids]
    neg_ratio = len(neg_ids) / len(questions)
    # 套卷规则：否定设问“应偏少”，默认阈值 0.25
    threshold = 0.25
    passed = neg_ratio <= threshold
    basis = str(data.get("semantic_basis", "") or "").strip()
    detail = (
        f"否定设问数量={len(neg_ids)}/{len(questions)}（占比={neg_ratio:.2f}，阈值<={threshold:.2f}）；"
        f"命中题号={neg_ids if neg_ids else '无'}"
    )
    if basis:
        detail += f"；语义依据={basis}"
    return {
        "name": "否定设问占比（LLM语义）",
        "status": "PASS" if passed else "ALERT",
        "detail": detail,
    }


def _mock_llm_response(input_value):
    text = str(input_value)
    if "QUALITY_SCORE" in text or "quality_score" in text:
        return "<QUALITY_SCORE>8.8</QUALITY_SCORE>"
    if "solver_validation" in text or "Cognitive Auditor" in text:
        return '{"solver_validation": {"passed": true, "predicted_answer": "A", "reasoning_path": "可唯一推出", "ambiguity_found": false}, "semantic_drift": {"passed": true, "evidence_quotes": ["证据句"], "drift_issues": []}}'
    if "distractor_quality" in text or "Value Assessor" in text:
        return '{"distractor_quality": {"score": 4, "homogeneity_issues": [], "weakness_issues": []}, "pedagogical_value": {"cognitive_level": "应用", "estimated_pass_rate": 0.62, "teaching_issues": []}, "risk_assessment": {"risk_level": "LOW", "risk_issues": []}}'
    if "generate_possible_answers(context)" in text or "只输出纯 Python 代码" in text:
        return (
            "def generate_possible_answers(context):\n"
            "    return [\n"
            "        {'type': 'correct', 'value': 100.0},\n"
            "        {'type': 'error_used_wrong_tax_rate_3_percent', 'value': 90.0},\n"
            "        {'type': 'error_forgot_to_deduct_vat', 'value': 80.0},\n"
            "    ]\n"
        )
    return '{"passed": true, "issues": []}'


def _to_str(v: Any) -> str:
    if v is None:
        return ""
    if isinstance(v, float) and math.isnan(v):
        return ""
    return str(v).strip()


def _infer_question_type(answer: str, options: list[str] | None = None) -> str:
    a = (answer or "").upper().strip()

    # 1) 直接口径：答案字段为“正确/错误”
    if "正确" in a or "错误" in a:
        return "true_false"

    # 2) 兼容口径：判断题答案写成 A/B，且前两项语义为正确/错误
    if a in {"A", "B"} and options and len(options) >= 2:
        a_opt = _to_str(options[0])
        b_opt = _to_str(options[1])
        if ("正确" in a_opt and "错误" in b_opt) or ("错误" in a_opt and "正确" in b_opt):
            return "true_false"

    # 3) 多选
    if "," in a or "、" in a or re.search(r"[A-D]\s*[，,、]\s*[A-D]", a):
        return "multiple_choice"

    # 4) 默认单选
    return "single_choice"


def _infer_assessment_type(row: dict[str, Any], stem: str, is_calculation: bool) -> str:
    label = _to_str(row.get("题目类型标签", ""))
    if label in {"基础概念/理解记忆", "实战应用/推演"}:
        return label

    if is_calculation:
        return "实战应用/推演"
    return "基础概念/理解记忆"


def _infer_is_calculation(row: dict[str, Any]) -> bool:
    """仅按表格显式标记判定是否计算题，不做关键词猜测。"""
    candidate_cols = [
        "是否计算题",
        "计算题标记",
        "题目计算题标签",
        "题目计算标签",
        "题型标记",
        "is_calculation",
    ]
    true_values = {"是", "y", "yes", "true", "1", "计算题"}
    for col in candidate_cols:
        raw = _to_str(row.get(col, ""))
        if not raw:
            continue
        if raw.strip().lower() in true_values:
            return True
        return False
    return False


def _parse_list_like(value: Any) -> list[str]:
    if isinstance(value, list):
        return [str(x).strip() for x in value if str(x).strip()]
    txt = _to_str(value)
    if not txt:
        return []
    if txt.startswith("[") and txt.endswith("]"):
        try:
            arr = json.loads(txt)
            if isinstance(arr, list):
                return [str(x).strip() for x in arr if str(x).strip()]
        except Exception:
            pass
    parts = re.split(r"[;\n；,，|]+", txt)
    return [p.strip() for p in parts if p.strip()]


def _load_questions_from_excel(path: Path, sheet_name: str, header_row: int) -> list[QuestionInput]:
    df = pd.read_excel(path, sheet_name=sheet_name, header=header_row)
    # 兼容“列名均为 Unnamed，首行才是业务表头”的 Excel 模板
    if list(df.columns) and all(str(c).startswith("Unnamed:") for c in df.columns):
        if len(df) > 0:
            header_values = [str(x).strip() for x in df.iloc[0].tolist()]
            if any(h in header_values for h in ["题目序号", "题干(必填)", "答案选项(必填)"]):
                df.columns = header_values
                df = df.iloc[1:].reset_index(drop=True)
    df = df.fillna("")

    items: list[QuestionInput] = []
    for idx, row in df.iterrows():
        stem = _to_str(row.get("题干(必填)", ""))
        if not stem:
            continue

        qid = _to_str(row.get("题目序号", "")) or str(idx + 1)
        options = [
            _to_str(row.get("选项A(必填)", "")),
            _to_str(row.get("选项B(必填)", "")),
            _to_str(row.get("选项C", "")),
            _to_str(row.get("选项D", "")),
        ]
        options = [x for x in options if x]
        answer = _to_str(row.get("答案选项(必填)", ""))
        explanation = _to_str(row.get("题目解析", ""))
        textbook_slice = _to_str(row.get("切片原文", ""))
        related_slices = (
            _parse_list_like(row.get("关联切片", ""))
            or _parse_list_like(row.get("关联切片原文", ""))
        )
        reference_slices = (
            _parse_list_like(row.get("参考切片", ""))
            or _parse_list_like(row.get("参考切片原文", ""))
        )
        mother_question = (
            _to_str(row.get("母题", ""))
            or _to_str(row.get("母题题干", ""))
            or _to_str(row.get("关联母题", ""))
        )
        term_locks = (
            _parse_list_like(row.get("术语锁词", ""))
            or _parse_list_like(row.get("锁词", ""))
        )
        mastery = _to_str(row.get("掌握程度", "")) or "未知"

        qtype = _infer_question_type(answer, options)

        is_calc = _infer_is_calculation(row)
        item = QuestionInput(
            question_id=f"EXCEL-{qid}",
            stem=stem,
            options=options,
            correct_answer=answer,
            explanation=explanation,
            textbook_slice=textbook_slice,
            related_slices=related_slices,
            reference_slices=reference_slices,
            mother_question=mother_question,
            term_locks=term_locks,
            mastery=mastery,
            question_type=qtype,
            is_calculation=is_calc,
            assessment_type=_infer_assessment_type(row, stem, is_calc),
        )
        items.append(item)

    return items


def _build_word_report(
    reports: list[dict],
    output_docx: Path,
    source_excel: Path,
    questions: list[QuestionInput],
    batch_checks: list[dict[str, str]] | None = None,
) -> None:
    def _reason_category(text: str) -> str:
        t = text
        # Strip leading tags like 【...】 to improve matching.
        while t.startswith("【") and "】" in t:
            t = t.split("】", 1)[1].strip()
        # Category heuristics.
        if any(k in t for k in ["格式", "括号", "标点", "单引号", "表格", "图片", "字数", "选项", "答案字段", "设问", "题干括号", "括号位置"]):
            return "格式"
        if any(k in t for k in ["知识", "超纲", "教材", "切片", "限定词", "知识边界", "证据链", "一致性"]):
            return "知识"
        if any(k in t for k in ["解析", "三段", "教材原文", "试题分析", "结论", "本题答案", "答案为"]):
            return "解析"
        if any(k in t for k in ["教学", "教学价值", "区分度", "通过率", "考核价值"]):
            return "教学"
        return "其他"

    def _normalize_reason(text: str) -> str:
        t = str(text or "").strip()
        # Drop leading bracketed tags.
        while t.startswith("【") and "】" in t:
            t = t.split("】", 1)[1].strip()
        # Normalize whitespace and common punctuation.
        t = re.sub(r"\s+", " ", t)
        t = t.replace("：", ":").strip()
        return t

    def _group_reasons(reasons: list[str]) -> list[tuple[str, list[str]]]:
        order = ["格式", "知识", "解析", "教学", "其他"]
        buckets: dict[str, list[str]] = {k: [] for k in order}
        seen: dict[str, set[str]] = {k: set() for k in order}
        for r in reasons:
            txt = str(r).strip()
            if not txt:
                continue
            cat = _reason_category(txt)
            norm = _normalize_reason(txt)
            if norm in seen[cat]:
                continue
            seen[cat].add(norm)
            buckets[cat].append(norm)
        return [(k, buckets[k]) for k in order if buckets[k]]

    def _top_reasons(items: list[dict], n: int = 8) -> list[tuple[str, int]]:
        c: Counter[str] = Counter()
        for it in items:
            for r in (it.get("reasons") or []):
                txt = str(r).strip()
                if not txt:
                    continue
                cat = _reason_category(txt)
                norm = _normalize_reason(txt)
                if norm:
                    c[f"{cat} | {norm}"] += 1
        return c.most_common(n)

    def _integrated_summary(it: dict) -> str:
        decision = str(it.get("decision", "unknown"))
        hard_pass = bool(it.get("hard_pass", False))
        scores = it.get("scores", {}) or {}
        logic = scores.get("logic", 0)
        knowledge = scores.get("knowledge", 0)
        distractor = scores.get("distractor", 0)
        teaching = scores.get("teaching", 0)
        risk = scores.get("risk", 0)
        reasons = [str(x) for x in (it.get("reasons") or []) if str(x).strip()]
        grouped = _group_reasons(reasons)

        dims = [
            ("logic", logic),
            ("knowledge", knowledge),
            ("distractor", distractor),
            ("teaching", teaching),
            ("risk", risk),
        ]
        dims_sorted = sorted(dims, key=lambda x: x[1])
        weak_dims = [k for k, v in dims_sorted if v <= 6][:2]

        if decision.endswith("PASS") or decision == "pass":
            verdict = "可通过：整体满足要求。"
        elif decision.endswith("REVIEW") or decision == "review":
            verdict = "需小修：存在可修复问题，建议修改后复核。"
        else:
            verdict = "不通过：当前不满足入库标准。"

        if not grouped:
            cause = "未发现明确问题。"
        else:
            flattened: list[str] = []
            for cat, items in grouped:
                for item in items:
                    flattened.append(f"{cat}:{item}")
            cause = "；".join(flattened[:3])

        if not hard_pass:
            action = "优先处理硬门禁问题（格式/表达/可解性），再处理质量优化。"
        elif weak_dims:
            action = f"优先提升维度：{', '.join(weak_dims)}。"
        else:
            action = "建议仅做小幅润色。"

        return f"结论：{verdict}\n归因：{cause}\n建议：{action}"

    doc = Document()
    doc.add_heading("离线Judge Excel评测报告", level=1)
    p = doc.add_paragraph()
    p.add_run("源文件：").bold = True
    p.add_run(str(source_excel))
    p = doc.add_paragraph()
    p.add_run("生成时间：").bold = True
    p.add_run(datetime.now().strftime("%Y-%m-%d %H:%M:%S"))

    action_dist = Counter(r.get("decision", "unknown") for r in reports)
    if batch_checks is None:
        batch_checks = _compute_batch_rule_checks(questions)

    doc.add_heading("汇总", level=2)
    doc.add_paragraph(f"总题数：{len(reports)}")
    doc.add_paragraph("Decision 分布：" + json.dumps(dict(action_dist), ensure_ascii=False))
    if batch_checks:
        doc.add_paragraph("套卷规则判定：")
        for c in batch_checks:
            doc.add_paragraph(
                f"[{c['status']}] {c['name']} - {c['detail']}",
                style="List Bullet",
            )

    top = _top_reasons(reports, n=8)
    if top:
        doc.add_paragraph("高频问题 Top：")
        for reason, cnt in top:
            doc.add_paragraph(f"{cnt} 次 - {reason}", style="List Bullet")

    doc.add_heading("逐题结果", level=2)
    for r in reports:
        doc.add_heading(r.get("question_id", ""), level=3)
        doc.add_paragraph(f"assessment_type: {r.get('assessment_type', '')}")
        doc.add_paragraph(f"decision: {r.get('decision')} | hard_pass: {r.get('hard_pass')} | confidence: {r.get('scores', {}).get('confidence')}")
        scores = r.get("scores", {})
        doc.add_paragraph(
            f"scores => logic={scores.get('logic')}, knowledge={scores.get('knowledge')}, distractor={scores.get('distractor')}, teaching={scores.get('teaching')}, risk={scores.get('risk')}"
        )
        reasons = [str(x) for x in (r.get("reasons") or []) if str(x).strip()]
        grouped = _group_reasons(reasons)
        doc.add_paragraph("reasons（根因去重）：")
        if not grouped:
            doc.add_paragraph("无", style="List Bullet")
        else:
            for cat, items in grouped:
                doc.add_paragraph(f"{cat}：{'; '.join(items)}", style="List Bullet")
        dim_results = r.get("dimension_results", {}) or {}
        if dim_results:
            doc.add_paragraph("分维执行结果：")
            for dim_name, dim_val in dim_results.items():
                status = str((dim_val or {}).get("status", "UNKNOWN"))
                issues = (dim_val or {}).get("issues", []) or []
                if issues:
                    doc.add_paragraph(f"[{status}] {dim_name}：{'; '.join(str(x) for x in issues)}", style="List Bullet")
                else:
                    default_msg = "未检测" if (status == "SKIP" or "SKIP" in str(status)) else "无问题"
                    doc.add_paragraph(f"[{status}] {dim_name}：{default_msg}", style="List Bullet")
                details = (dim_val or {}).get("details", {}) or {}
                unsupported = details.get("unsupported_options", []) or []
                why_unrelated = details.get("why_unrelated", []) or []
                overlap_pairs = details.get("overlap_pairs", []) or []
                stem_option_conflicts = details.get("stem_option_conflicts", []) or []
                mutual_exclusivity_fail = details.get("mutual_exclusivity_fail", False)
                if unsupported or why_unrelated:
                    doc.add_paragraph(
                        f"选项级证据 -> unsupported_options={unsupported}; why_unrelated={why_unrelated}",
                        style="List Bullet",
                    )
                if overlap_pairs:
                    doc.add_paragraph(
                        f"选项重叠证据 -> overlap_pairs={overlap_pairs}",
                        style="List Bullet",
                    )
                if stem_option_conflicts:
                    doc.add_paragraph(
                        f"题干-选项冲突证据 -> stem_option_conflicts={stem_option_conflicts}",
                        style="List Bullet",
                    )
                if mutual_exclusivity_fail:
                    doc.add_paragraph(
                        "互斥性证据 -> mutual_exclusivity_fail=true",
                        style="List Bullet",
                    )
                term_mismatch_issues = details.get("term_mismatch_issues", []) or []
                if term_mismatch_issues:
                    doc.add_paragraph("术语问题（原词->建议词->位置）：", style="List Bullet")
                    for item in term_mismatch_issues:
                        if isinstance(item, dict):
                            raw_term = str(item.get("raw_term", "")).strip()
                            suggested = str(item.get("suggested_term", "")).strip()
                            location = str(item.get("location", "")).strip()
                            source = str(item.get("source", "")).strip()
                            doc.add_paragraph(
                                f"{raw_term} -> {suggested} -> {location}" + (f" ({source})" if source else ""),
                                style="List Bullet",
                            )
        doc.add_paragraph("整合说明：")
        doc.add_paragraph(_integrated_summary(r))

        obs = r.get("observability", {})
        tok = (obs.get("tokens") or {})
        doc.add_paragraph(
            f"observability => llm_calls={obs.get('llm_calls',0)}, failed_calls={obs.get('failed_calls',0)}, latency_ms={obs.get('latency_ms',0)}, prompt_tokens={tok.get('prompt_tokens',0)}, completion_tokens={tok.get('completion_tokens',0)}"
        )
        if obs.get("last_error"):
            doc.add_paragraph(f"last_error => {obs.get('last_error')}")

        costs = r.get("costs", {})
        doc.add_paragraph(f"costs => per_question_usd={costs.get('per_question_usd',0)}")

    output_docx.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_docx)


def main():
    parser = argparse.ArgumentParser(description="Run Offline Judge on Excel and export Word report")
    parser.add_argument("--input-excel", type=Path, required=True)
    parser.add_argument("--output-docx", type=Path, required=True)
    parser.add_argument("--output-json", type=Path, default=None)
    parser.add_argument("--sheet", default="AI题目")
    parser.add_argument("--header-row", type=int, default=1, help="0-based header row index in excel")
    parser.add_argument("--provider", choices=["openai", "anthropic", "ait"], default="openai")
    parser.add_argument("--model", default=None)
    parser.add_argument("--temperature", type=float, default=0)
    parser.add_argument("--limit", type=int, default=0, help="仅评测前N题，0表示全部")
    parser.add_argument("--mock-llm", action="store_true")
    parser.add_argument(
        "--progress-json",
        type=Path,
        default=None,
        help="实时进度文件（每题更新一次）",
    )
    args = parser.parse_args()

    load_dotenv()

    questions = _load_questions_from_excel(args.input_excel, args.sheet, args.header_row)
    if args.limit and args.limit > 0:
        questions = questions[: args.limit]
    if args.mock_llm:
        llm = RunnableLambda(_mock_llm_response)
    else:
        llm = build_llm(provider=args.provider, model=args.model, temperature=args.temperature)

    reports: list[dict] = []
    total = len(questions)

    progress_path: Path | None = args.progress_json
    if progress_path is None and args.output_json:
        progress_path = args.output_json.with_name(args.output_json.stem + ".progress.json")
    elif progress_path is None:
        progress_path = args.output_docx.with_suffix(".progress.json")

    for i, q in enumerate(questions, start=1):
        r = run_judge(q, llm).model_dump()
        reports.append(r)

        # 实时进度输出（便于终端观察）
        decision = str(r.get("decision", "unknown"))
        obs = r.get("observability", {}) or {}
        print(
            f"[{i}/{total}] {q.question_id} decision={decision} llm_calls={obs.get('llm_calls', 0)} failed_calls={obs.get('failed_calls', 0)}",
            flush=True,
        )

        # 实时落盘（便于外部查看进度）
        progress_payload = {
            "total": total,
            "completed": i,
            "updated_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            "latest_question_id": q.question_id,
            "latest_decision": decision,
            "reports": reports,
        }
        progress_path.parent.mkdir(parents=True, exist_ok=True)
        progress_path.write_text(
            json.dumps(progress_payload, ensure_ascii=False, indent=2),
            encoding="utf-8",
        )

    if args.output_json:
        args.output_json.parent.mkdir(parents=True, exist_ok=True)
        args.output_json.write_text(json.dumps(reports, ensure_ascii=False, indent=2), encoding="utf-8")

    batch_checks = _compute_batch_rule_checks(questions)
    llm_negation_check = _compute_batch_llm_negation_check(questions, llm)
    if llm_negation_check is not None:
        batch_checks.append(llm_negation_check)

    _build_word_report(
        reports,
        args.output_docx,
        args.input_excel,
        questions,
        batch_checks=batch_checks,
    )

    print(f"questions: {len(questions)}")
    print(f"docx: {args.output_docx}")
    if args.output_json:
        print(f"json: {args.output_json}")
    print(f"progress_json: {progress_path}")


if __name__ == "__main__":
    main()
