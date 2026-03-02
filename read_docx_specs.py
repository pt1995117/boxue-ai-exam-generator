
import docx
import os

docx_path = '/Users/panting/Desktop/搏学考试/AI出题/最终完整的切片样式规范.docx'

if not os.path.exists(docx_path):
    print(f"File not found: {docx_path}")
    exit(1)

doc = docx.Document(docx_path)

print("### Document Content ###")
for para in doc.paragraphs:
    if para.text.strip():
        print(para.text)

print("\n### Tables ###")
for table in doc.tables:
    for row in table.rows:
        row_text = []
        for cell in row.cells:
            row_text.append(cell.text.strip())
        print(" | ".join(row_text))
    print("--- Table End ---")
